import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class ResumeFormatter:
    def __init__(self):
        """Initialize the resume formatter."""
        pass

    def modify_resume(self, resume_data, parsed_job_description):
        """
        Modify resume data based on job description requirements.
        
        Args:
            resume_data (dict): Original resume data
            parsed_job_description (list): List of keywords from job description
            
        Returns:
            dict: Modified resume data
        """
        try:
            modified_data = resume_data.copy()
            
            # Modify summary to include relevant keywords
            modified_data["summary"] = self._enhance_summary(
                resume_data.get("summary", ""), 
                parsed_job_description
            )
            
            # Enhance skills section
            modified_data["skills"] = self._enhance_skills(
                resume_data.get("skills", []), 
                parsed_job_description
            )
            
            # Enhance experience descriptions
            modified_data["experience"] = self._enhance_experience(
                resume_data.get("experience", []), 
                parsed_job_description
            )
            
            # Enhance project descriptions
            modified_data["projects"] = self._enhance_projects(
                resume_data.get("projects", []), 
                parsed_job_description
            )
            
            # Store job requirements for reference
            modified_data["job_requirements"] = parsed_job_description
            
            logging.debug(f"Modified resume data: {modified_data}")
            return modified_data
            
        except Exception as e:
            logging.error(f"Error modifying resume: {e}")
            return resume_data

    def _enhance_summary(self, original_summary, job_requirements):
        """Enhance the professional summary with relevant keywords."""
        if not original_summary:
            original_summary = "Experienced professional with strong technical skills"
        
        # Extract the most relevant keywords for summary
        key_skills = [req for req in job_requirements if not any(char.isdigit() for char in req)][:5]
        
        if key_skills:
            # Add relevant skills to summary
            enhanced_summary = f"{original_summary.rstrip('.')}. Experienced in {', '.join(key_skills[:3])}"
            if len(key_skills) > 3:
                enhanced_summary += f" and {key_skills[3]}"
            enhanced_summary += "."
        else:
            enhanced_summary = original_summary
        
        return enhanced_summary

    def _enhance_skills(self, original_skills, job_requirements):
        """Enhance skills section with relevant job requirements."""
        enhanced_skills = original_skills.copy()
        
        # Add job requirements that aren't already in skills
        for requirement in job_requirements:
            # Skip experience requirements (containing numbers/years)
            if any(char.isdigit() for char in requirement):
                continue
                
            # Check if skill is already present (case-insensitive)
            if not any(requirement.lower() in skill.lower() for skill in enhanced_skills):
                enhanced_skills.append(requirement)
        
        # Prioritize skills that match job requirements
        matched_skills = []
        other_skills = []
        
        for skill in enhanced_skills:
            if any(req.lower() in skill.lower() or skill.lower() in req.lower() 
                   for req in job_requirements):
                matched_skills.append(skill)
            else:
                other_skills.append(skill)
        
        # Return matched skills first, then others
        return matched_skills + other_skills

    def _enhance_experience(self, original_experience, job_requirements):
        """Enhance work experience with relevant keywords."""
        enhanced_experience = []
        
        # Technical keywords to add to experience
        tech_keywords = [req for req in job_requirements 
                        if not any(char.isdigit() for char in req)]
        
        for exp in original_experience:
            enhanced_exp = exp
            
            # Add relevant keywords to experience descriptions
            if tech_keywords and len(exp) > 50:  # Only enhance substantial entries
                # Add 1-2 relevant keywords to experience
                relevant_keywords = tech_keywords[:2]
                for keyword in relevant_keywords:
                    if keyword.lower() not in exp.lower():
                        enhanced_exp += f" Utilized {keyword} for improved performance."
                        break
            
            enhanced_experience.append(enhanced_exp)
        
        return enhanced_experience

    def _enhance_projects(self, original_projects, job_requirements):
        """Enhance project descriptions with relevant technologies."""
        enhanced_projects = []
        
        # Technical keywords for projects
        tech_keywords = [req for req in job_requirements 
                        if not any(char.isdigit() for char in req)]
        
        for project in original_projects:
            enhanced_project = project
            
            # Add relevant technologies to project descriptions
            if tech_keywords and len(project) > 30:
                relevant_tech = tech_keywords[:1]  # Add one relevant technology
                for tech in relevant_tech:
                    if tech.lower() not in project.lower():
                        enhanced_project += f" Implemented using {tech}."
                        break
            
            enhanced_projects.append(enhanced_project)
        
        # Add a new project if original list is short and we have relevant requirements
        if len(enhanced_projects) < 2 and tech_keywords:
            new_project = f"Personal project demonstrating proficiency in {tech_keywords[0]}"
            if len(tech_keywords) > 1:
                new_project += f" and {tech_keywords[1]}"
            enhanced_projects.append(new_project)
        
        return enhanced_projects

    def create_enhanced_document(self, resume_data, output_path, template_id=1):
        """
        Create an enhanced resume document with professional formatting.
        
        Args:
            resume_data (dict): Enhanced resume data
            output_path (str): Path to save the output document
            template_id (int): Template ID to use for formatting
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Create new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add content based on template
            self._add_header(doc, resume_data)
            self._add_summary(doc, resume_data)
            self._add_skills(doc, resume_data)
            self._add_experience(doc, resume_data)
            self._add_projects(doc, resume_data)
            self._add_education(doc, resume_data)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating document: {str(e)}")
            return False
    
    def _add_header(self, doc, resume_data):
        """Add header section with contact information"""
        contact = resume_data.get('contact', {})
        
        # Name (Title)
        name_paragraph = doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(contact.get('name', 'Professional Resume'))
        name_run.font.size = Pt(24)
        name_run.bold = True
        
        # Contact information
        if contact:
            contact_paragraph = doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_info = []
            if contact.get('email'):
                contact_info.append(contact['email'])
            if contact.get('phone'):
                contact_info.append(contact['phone'])
            if contact.get('location'):
                contact_info.append(contact['location'])
            
            contact_run = contact_paragraph.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(12)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_summary(self, doc, resume_data):
        """Add professional summary section"""
        summary = resume_data.get('summary', '')
        if summary:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('PROFESSIONAL SUMMARY')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            summary_paragraph = doc.add_paragraph(summary)
            summary_paragraph.style = 'Normal'
            
            doc.add_paragraph()
    
    def _add_skills(self, doc, resume_data):
        """Add skills section"""
        skills = resume_data.get('skills', [])
        if skills:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('CORE COMPETENCIES')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            # Format skills in columns
            skills_text = ' • '.join(skills)
            skills_paragraph = doc.add_paragraph(skills_text)
            
            doc.add_paragraph()
    
    def _add_experience(self, doc, resume_data):
        """Add work experience section"""
        experience = resume_data.get('experience', [])
        if experience:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('PROFESSIONAL EXPERIENCE')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            for exp in experience:
                # Job title and company
                job_paragraph = doc.add_paragraph()
                title_run = job_paragraph.add_run(exp.get('title', 'Position'))
                title_run.bold = True
                title_run.font.size = Pt(12)
                
                company_run = job_paragraph.add_run(f" | {exp.get('company', 'Company')}")
                company_run.font.size = Pt(12)
                
                # Dates
                if exp.get('dates'):
                    dates_paragraph = doc.add_paragraph(exp['dates'])
                    dates_paragraph.style = 'Normal'
                
                # Description
                if exp.get('description'):
                    desc_paragraph = doc.add_paragraph(exp['description'])
                    desc_paragraph.style = 'Normal'
                
                doc.add_paragraph()
    
    def _add_projects(self, doc, resume_data):
        """Add projects section"""
        projects = resume_data.get('projects', [])
        if projects:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('KEY PROJECTS')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            for project in projects:
                # Project title
                project_paragraph = doc.add_paragraph()
                title_run = project_paragraph.add_run(project.get('title', 'Project'))
                title_run.bold = True
                title_run.font.size = Pt(12)
                
                # Description
                if project.get('description'):
                    desc_paragraph = doc.add_paragraph(project['description'])
                    desc_paragraph.style = 'Normal'
                
                doc.add_paragraph()
    
    def _add_education(self, doc, resume_data):
        """Add education section"""
        education = resume_data.get('education', [])
        if education:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('EDUCATION')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            
            for edu in education:
                edu_paragraph = doc.add_paragraph()
                degree_run = edu_paragraph.add_run(edu.get('degree', 'Degree'))
                degree_run.bold = True
                
                if edu.get('school'):
                    school_run = edu_paragraph.add_run(f" | {edu['school']}")
                
                if edu.get('year'):
                    year_paragraph = doc.add_paragraph(edu['year'])
    
    def add_content(self, doc, resume_data):
        """
        Add formatted content to the Word document.
        
        Args:
            doc (Document): Python-docx Document object
            resume_data (dict): Resume data to format
        """
        try:
            # Document title
            title = doc.add_heading('Professional Resume', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact Information
            if resume_data.get("contact_info"):
                contact_para = doc.add_paragraph()
                contact_info = resume_data["contact_info"]
                contact_text = []
                if contact_info.get("email"):
                    contact_text.append(f"Email: {contact_info['email']}")
                if contact_info.get("phone"):
                    contact_text.append(f"Phone: {contact_info['phone']}")
                if contact_text:
                    contact_para.add_run(" | ".join(contact_text))
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()  # Add spacing
            
            # Professional Summary
            if resume_data.get("summary"):
                doc.add_heading('Professional Summary', level=1)
                summary_para = doc.add_paragraph(resume_data["summary"])
                doc.add_paragraph()  # Add spacing
            
            # Skills
            if resume_data.get("skills"):
                doc.add_heading('Technical Skills', level=1)
                skills_text = " • ".join(resume_data["skills"])
                doc.add_paragraph(skills_text)
                doc.add_paragraph()  # Add spacing
            
            # Work Experience
            if resume_data.get("experience"):
                doc.add_heading('Professional Experience', level=1)
                for exp in resume_data["experience"]:
                    exp_para = doc.add_paragraph()
                    exp_para.add_run("• ").bold = True
                    exp_para.add_run(exp)
                doc.add_paragraph()  # Add spacing
            
            # Projects
            if resume_data.get("projects"):
                doc.add_heading('Notable Projects', level=1)
                for project in resume_data["projects"]:
                    project_para = doc.add_paragraph()
                    project_para.add_run("• ").bold = True
                    project_para.add_run(project)
                doc.add_paragraph()  # Add spacing
            
            # Education
            if resume_data.get("education"):
                doc.add_heading('Education', level=1)
                for edu in resume_data["education"]:
                    edu_para = doc.add_paragraph()
                    edu_para.add_run("• ").bold = True
                    edu_para.add_run(edu)
            
            logging.debug("Content added to document successfully")
            
        except Exception as e:
            logging.error(f"Error adding content to document: {e}")
            # Add fallback content
            doc.add_heading('Resume', 0)
            doc.add_paragraph("Resume content could not be properly formatted.")
            if resume_data.get("summary"):
                doc.add_paragraph(f"Summary: {resume_data['summary']}")
            if resume_data.get("skills"):
                doc.add_paragraph(f"Skills: {', '.join(resume_data['skills'])}")
